import streamlit as st
from PIL import Image
import pandas as pd
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
import pytesseract  # Asegúrate de tener instalada esta biblioteca (pip install pytesseract)
import os
from dotenv import load_dotenv, find_dotenv
from openai import OpenAI
from pydub import AudioSegment
from pydub.playback import play
import time

import openai
from tenacity import retry

# Configuración de Tesseract (ajusta según tu sistema operativo)
           #pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
#pytesseract.pytesseract.tesseract_cmd = r"tesseract.exe"

#current_dir = os.path.dirname(os.path.abspath(__file__))  # Obtiene el directorio actual
#pytesseract.pytesseract.tesseract_cmd = os.path.join(current_dir, "tesseract.exe")  # Establece el path dinámico

# Configuración de Tesseract (ajusta según tu sistema operativo o entorno)
if os.name == 'nt':  # Si está en Windows
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
else:  # Si está en un sistema Linux o basado en la nube
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"  # Ruta típica en Linux

# Verifica que Tesseract esté instalado correctamente
try:
    tesseract_version = pytesseract.get_tesseract_version()
    print(f"Tesseract OCR versión detectada: {tesseract_version}")
except Exception as e:
    print(f"Error detectando Tesseract OCR: {e}")
    st.error("Tesseract OCR no está instalado o configurado correctamente. Verifica la configuración.")

# Configuración de OpenAI para la generación de audio y procesamiento de texto
load_dotenv(find_dotenv(), override=True)  # Carga las variables de entorno desde un archivo .env

# Obtiene la clave de API de Open

# Configuración de OpenAI para la generación de audio
load_dotenv(find_dotenv(), override=True)

apikey = os.environ.get('OPENAI_API_KEY')

client = OpenAI(api_key=apikey)

# Configuración de la clave de API de OpenAI

openai.api_key = os.getenv("OPENAI_API_KEY")  # Cambia a tu clave directamente si no usas .env


# Configuración directa de la clave de API
#apikey = 'aki va la ip'

# Configurar el cliente OpenAI con la clave
#client = OpenAI(api_key=apikey)

# También configurar `openai.api_key` (opcional, pero usualmente necesario)
#openai.api_key = apikey

# Funciones auxiliares
@st.cache_data
def cargar_imagen(imagen_file):
    img = Image.open(imagen_file)
    return img

def leer_pdf(file):
    pdf_reader = PdfReader(file)
    todo_el_texto = ""
    for pagina in pdf_reader.pages:
        todo_el_texto += pagina.extract_text()
    return todo_el_texto

def crear_documento_word(texto, nombre_archivo="transcripcion.docx"):
    doc = Document()
    doc.add_heading("Texto Extraído", level=1)
    doc.add_paragraph(texto)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, nombre_archivo

def extraer_texto_de_imagen(imagen, idioma="eng"):
    try:
        texto = pytesseract.image_to_string(imagen, lang=idioma)
        return texto
    except Exception as e:
        return f"Error al procesar la imagen: {str(e)}"


def generar_audio_con_openai(text, voice):
    try:
        response = client.audio.speech.create(
            model="tts-1",
            voice=voice,
            input=text,
        )
        audio_path = "audio.mp3"
        with open(audio_path, "wb") as output_file:
            for chunk in response.iter_bytes():
                if chunk:
                    output_file.write(chunk)
        return audio_path
    except Exception as e:
        return None, str(e)


# Función para transcribir audio a texto
def transcribir_audio(audio_file):
    try:
        transcription = client.audio.transcriptions.create(
            file=audio_file,
            model="whisper-1",
            response_format="verbose_json",
            language="es"
        )
        return transcription.text
    except Exception as e:
        return f"Error al transcribir el audio: {str(e)}"

# Nueva opción del menú para cargar audio y transcribir
def cargar_y_transcribir_audio():
    st.subheader("Cargar Audio y Transcribir a Texto")
    audio_file = st.file_uploader("Carga un archivo de audio (mp3, wav)", type=["mp3", "wav"])

    if audio_file is not None:
        # Mostrar detalles del archivo
        detalle_archivo = {
            "nombre_archivo": audio_file.name,
            "tipo_archivo": audio_file.type,
            "tamaño_archivo": audio_file.size,
        }
        st.write(detalle_archivo)

        # Reproducir audio
        st.audio(audio_file, format="audio/wav")

        # Transcribir audio
        if st.button("Transcribir Audio"):
            with st.spinner("Transcribiendo el audio..."):
                texto_transcrito = transcribir_audio(audio_file)
                if texto_transcrito:
                    st.text_area("Texto Transcrito", texto_transcrito, height=300)
                else:
                    st.warning("No se pudo transcribir el audio. Verifique que el archivo sea válido y audible.")


# Menú Principal
def main():
    st.title("Carga de Archivos y Generación de Audio")

    # Opciones del menú
    menu = [
        "Pasar de Imagen a Texto",
        "Conjunto de Datos",
        "Archivos de Documentos",
        "Generación de Audio desde texto",
        "Cargar Audio y Transcribir a Texto"
        ]

    election = st.sidebar.radio("Selecciona una opción del Menú", menu)

    if election == "Pasar de Imagen a Texto":
        #st.subheader("Sección de Imágenes")
        imagen = st.file_uploader("Carga una imagen", type=["png", "jpg", "jpeg"])
        if imagen is not None:
            # Mostrar detalles del archivo
            detalle_archivo = {
                "nombre_archivo": imagen.name,
                "tipo_archivo": imagen.type,
                "tamaño_archivo": imagen.size,
            }
            st.write(detalle_archivo)

            # Mostrar la imagen cargada
            imagen_cargada = cargar_imagen(imagen)
            st.image(imagen_cargada, width=250, caption="Imagen cargada exitosamente.")

            # Botón para interpretar la imagen
            if st.button("Interpretar Imagen"):
                with st.spinner("Interpretando la imagen..."):
                    texto_extraido = extraer_texto_de_imagen(imagen_cargada)
                if texto_extraido.strip():
                    st.subheader("Texto Interpretado de la Imagen:")
                    st.text_area("Resultado", texto_extraido, height=300)
                else:
                    st.warning("No se pudo extraer texto de la imagen. Verifique que contiene texto legible.")

    elif election == "Conjunto de Datos":
        #st.subheader("Sección de Conjuntos de Datos")
        archivo_datos = st.file_uploader("Carga un archivo CSV o Excel", type=["csv", "xlsx"])
        if archivo_datos is not None:
            detalle_archivo = {
                "nombre_archivo": archivo_datos.name,
                "tipo_archivo": archivo_datos.type,
                "tamaño_archivo": archivo_datos.size,
            }
            st.write(detalle_archivo)

            if archivo_datos.name.endswith(".csv"):
                delimitador = st.text_input("Especifica el delimitador (por defecto: ',')", value=",")
                codificacion = st.selectbox("Selecciona la codificación del archivo", ["utf-8", "latin1", "ISO-8859-1"], index=0)
                if st.button("Cargar CSV"):
                    try:
                        df = pd.read_csv(archivo_datos, delimiter=delimitador, encoding=codificacion)
                        st.write("Vista previa de los datos:")
                        st.dataframe(df)
                    except Exception as e:
                        st.error(f"Error al cargar el archivo CSV: {str(e)}")

            elif archivo_datos.name.endswith(".xlsx"):
                if st.button("Cargar Excel"):
                    try:
                        df = pd.read_excel(archivo_datos)
                        st.write("Vista previa de los datos:")
                        st.dataframe(df)
                    except Exception as e:
                        st.error(f"Error al cargar el archivo Excel: {str(e)}")

    elif election == "Archivos de Documentos":
        #st.subheader("Sección de Documentos")
        archivo_pdf = st.file_uploader("Carga un archivo PDF", type=["pdf"])
        if archivo_pdf is not None:
            texto_extraido = leer_pdf(archivo_pdf)
            st.text_area("Contenido del PDF", texto_extraido, height=300)

            if st.button("Generar documento"):
                buffer, nombre_archivo = crear_documento_word(texto_extraido)
                st.download_button(
                    label="Descargar Archivo Word",
                    data=buffer,
                    file_name=nombre_archivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    elif election == "Generación de Audio desde texto":
        st.subheader("Generación de Audio")
        text = st.text_area("Ingrese el texto para generar el audio:", height=200)
        voices = ["alloy", "echo", "fable", "onyx", "nova", "shimmer"]
        voice = st.selectbox("Seleccione la voz:", voices)

          # boton para generar el audio
        if st.button("Generar el audio"):
            if text:
                response = client.audio.speech.create(
                    model="tts-1",
                    voice=voice,
                    input=text,

                )
                audio_path = "audio.mp3"
                with open(audio_path, "wb") as output_file:
                    for chunk in response.iter_bytes():
                        if chunk:
                            output_file.write(chunk)
                st.success(f"Audio generado y guardado en {audio_path}")

                audio_file = open(audio_path, 'rb')
                audio_bytes = audio_file.read()
                st.audio(audio_bytes, format='audio/mp3')
            else:
                st.error("Por favor, ingrese un texto")

    elif election == "Cargar Audio y Transcribir a Texto":
        cargar_y_transcribir_audio()

if __name__ == '__main__':
    main()
